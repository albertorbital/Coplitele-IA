#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import sys

# Data dictionary representing all website content in ES, CA, and EN
translations = {
    "es": {
        "menu_inicio": "Inicio",
        "menu_proyecto": "Proyecto",
        "menu_impacto": "Impacto y Difusión",
        "hero_tag": "Proyecto de Investigación activo",
        "hero_main_title": "Rediseñamos el futuro de la educación con <span class=\"blue-highlight\">Codiseño</span> e <span class=\"green-highlight\">Inteligencia Artificial</span>",
        "hero_desc": "COPLITELE-IA es un proyecto de investigación que transforma la educación superior integrando la Inteligencia Artificial Generativa (IAG) desde un enfoque pedagógico innovador. A través del codiseño educativo entre docentes y estudiantes, impulsamos la personalización del aprendizaje mediante itinerarios flexibles, promoviendo entornos virtuales conectados que garantizan la equidad, la inclusión y la calidad educativa.",
        "btn_conocer": "Conoce el Proyecto",
        "btn_publicaciones": "Publicaciones",
        "news_title": "Últimas noticias",
        "news_pretitle": "Actualidad",
        "stats_years": "Años de investigación",
        "stats_investigadores": "Investigadores",
        "stats_publicaciones": "Publicaciones",
        "stats_experiencias": "Experiencias",
        "progress_label": "Progreso del Proyecto",
        "submenu_desc": "Descripción",
        "submenu_obj": "Objetivos",
        "submenu_miembros": "Miembros",
        "submenu_noticias": "Últimas noticias",
        "submenu_transferencia": "Transferencia",
        "submenu_publicaciones": "Producción científica",
        "submenu_recursos": "Recursos",
        "obj_title": "Objetivos del Proyecto",
        "obj_1_title": "Objetivo Principal",
        "obj_1_desc": "Investigar el impacto de la IA en la telecolaboración interlingüística, analizando cómo el codiseño mejora los procesos de aprendizaje de lenguas.",
        "obj_2_title": "Innovación",
        "obj_2_desc": "Desarrollar herramientas de personalización adaptativa basadas en IA para entornos de aprendizaje de lenguas mediado por tecnología.",
        "obj_3_title": "Alcance Internacional",
        "obj_3_desc": "Colaboración con instituciones europeas y latinoamericanas para estudiar la comunicación mediada por ordenador en contextos plurilingüísticos.",
        "obj_4_title": "Impacto Social",
        "obj_4_desc": "Contribuir a la inclusión digital y educativa a través de soluciones tecnológicas accesibles y personalizadas.",
        "project_pretitle": "Conoce más",
        "project_title": "El Proyecto",
        "tab_desc": "Descripción",
        "tab_equipo": "Equipo",
        "project_what_is": "¿Qué es COPLITELE-IA?",
        "project_what_is_p1": "<strong>COPLITELE-IA</strong> (Codiseño, Personalización y Tecnología en la Telecolaboración para el Aprendizaje de Lenguas con IA) es un proyecto de investigación financiado por el Ministerio de Ciencia e Innovación de España (Ref. PID2021-127836NB-I00).",
        "project_what_is_p2": "El proyecto examina cómo los enfoques de codiseño participativo, combinados con herramientas de inteligencia artificial, pueden transformar las experiencias de telecolaboración interlingüística, especialmente en contextos de aprendizaje de lenguas extranjeras.",
        "project_what_is_p3": "Nuestro enfoque integra la comunicación mediada por ordenador (CMC), el aprendizaje de lenguas asistido por ordenador (CALL) y las tecnologías emergentes de IA, incluidos modelos de lenguaje, análisis de sentimientos y sistemas de retroalimentación automatizada.",
        "meta_ref": "Ref. Proyecto",
        "meta_duracion": "Duración",
        "meta_lider": "Institución Líder",
        "meta_financiacion": "Financiación",
        "framework_header": "COPLITELE-IA",
        "framework_subtitle": "Framework de investigación",
        "trans_pretitle": "Impacto y difusión",
        "trans_title": "Actividades de Transferencia",
        "trans_subtitle": "Actividades diseñadas para maximizar el impacto social y educativo de nuestra investigación, conectando la academia con la práctica docente y la sociedad.",
        "trans_btn_more": "Ver todas las actividades",
        "pub_pretitle": "Biblioteca Científica",
        "pub_title": "Producción Científica",
        "pub_subtitle": "Accede a las publicaciones del proyecto sincronizadas con nuestra biblioteca de referencias de Zotero.",
        "search_placeholder": "Buscar por título, autor o tag...",
        "tab_all": "Todas",
        "tab_revistas": "Revistas (Zotero)",
        "tab_libros": "Libros (Zotero)",
        "tab_ponencias": "Ponencias y Congresos",
        "rec_pretitle": "Repositorio Abierto",
        "rec_title": "Recursos del Proyecto",
        "rec_subtitle": "Descarga materiales pedagógicos y herramientas de codiseño desarrolladas durante la investigación.",
        "rec_btn_pdf": "Descargar PDF (4.8 MB)",
        "rec_btn_zip": "Descargar Plantillas (Zip, 12 MB)",
        "rec_btn_git": "Ver Repositorio GitHub",
        "colab_title": "Con la colaboración y financiación de:",
        "footer_copy": "&copy; 2026 COPLITELE-IA. Proyecto financiado por el Ministerio español de Ciencia e Innovación, desarrollado por el GTE de la UIB y el IRIE."
    },
    "ca": {
        "menu_inicio": "Inici",
        "menu_proyecto": "Projecte",
        "menu_impacto": "Impacte i Difusió",
        "hero_tag": "Proyecte d’Investigació actiu",
        "hero_main_title": "Redissenyam el futur de l'educació amb <span class=\"blue-highlight\">Codisseny</span> i <span class=\"green-highlight\">Intel·ligència Artificial</span>",
        "hero_desc": "COPLITELE-IA és un projecte d'investigació que transforma l'educació superior integrant la Intel·ligència Artificial Generativa (IAG) des d'un enfocament pedagògic innovador. A través del codissenyi educatiu entre docents i estudiants, impulsem la personalització de l'aprenentatge mitjançant itineraris flexibles, promovent entorns virtuals connectats que garanteixen l'equitat, la inclusió i la qualitat educativa.",
        "btn_conocer": "Conèix el Proyecte",
        "btn_publicaciones": "Publicacions",
        "news_title": "Últimes notícies",
        "news_pretitle": "Actualitat",
        "stats_years": "Anys d’investigació",
        "stats_investigadores": "Investigadors",
        "stats_publicaciones": "Publicacions",
        "stats_experiencias": "Experiències",
        "progress_label": "Progrés del Projecte",
        "submenu_desc": "Descripció",
        "submenu_obj": "Objectius",
        "submenu_miembros": "Membres",
        "submenu_noticias": "Últimes notícies",
        "submenu_transferencia": "Transferència",
        "submenu_publicaciones": "Producció científica",
        "submenu_recursos": "Recursos",
        "obj_title": "Objectius del Projecte",
        "obj_1_title": "Objectiu Principal",
        "obj_1_desc": "Investigar l'impacte de la IA en la telecolaboració interlingüística, analitzant com el codisseny millora els processos d'aprenentatge de llengües.",
        "obj_2_title": "Innovació",
        "obj_2_desc": "Desenvolupar eines de personalització adaptativa basades en IA per a entorns d'aprenentatge de llengües mediat per tecnologia.",
        "obj_3_title": "Abast Internacional",
        "obj_3_desc": "Col·laboració amb institucions europees i llatinoamericanes per estudiar la comunicació mediada per ordinador en contextos plurilingües.",
        "obj_4_title": "Impacte Social",
        "obj_4_desc": "Contribuir a la inclusió digital i educativa a través de solucions tecnològiques accessibles i personalitzades.",
        "project_pretitle": "Coneix-ne més",
        "project_title": "El Projecte",
        "tab_desc": "Descripció",
        "tab_equipo": "Equip",
        "project_what_is": "Què és COPLITELE-IA?",
        "project_what_is_p1": "<strong>COPLITELE-IA</strong> (Codisseny, Personalització i Tecnologia en la Telecolaboració per a l'Aprenentatge de Llengües amb IA) és un projecte d'investigació finançat pel Ministeri de Ciència i Innovació d'Espanya (Ref. PID2021-127836NB-I00).",
        "project_what_is_p2": "El projecte examina com els enfocaments de codisseny participatiu, combinats amb eines d'intel·ligència artificial, poden transformar les experiències de telecolaboració interlingüística, especialment en contextos d'aprenentatge de llengües estrangeres.",
        "project_what_is_p3": "El nostre enfocament integra la comunicació mediada per ordinador (CMC), l'aprenentatge de llengües assistit per ordinador (CALL) i les tecnologies emergents d'IA, inclosos models de llenguatge, anàlisi de sentiments i sistemes de retroalimentació automatizada.",
        "meta_ref": "Ref. Projecte",
        "meta_duracion": "Durada",
        "meta_lider": "Institució Líder",
        "meta_financiacion": "Finançament",
        "framework_header": "COPLITELE-IA",
        "framework_subtitle": "Framework d'investigació",
        "trans_pretitle": "Impacte i difusió",
        "trans_title": "Activitats de Transferència",
        "trans_subtitle": "Activitats dissenyades per maximitzar l'impacte social i educatiu de la nostra investigació, connectant l'acadèmia amb la pràctica docent i la societat.",
        "trans_btn_more": "Veure totes les activitats",
        "pub_pretitle": "Biblioteca Científica",
        "pub_title": "Producció Científica",
        "pub_subtitle": "Accedeix a les publicacions del projecte sincronitzades amb la nostra biblioteca de referències de Zotero.",
        "search_placeholder": "Cercar por títol, autor o etiqueta...",
        "tab_all": "Totes",
        "tab_revistas": "Revistes (Zotero)",
        "tab_libros": "Llibres (Zotero)",
        "tab_ponencias": "Ponències i Congressos",
        "rec_pretitle": "Repositori Obert",
        "rec_title": "Recursos del Projecte",
        "rec_subtitle": "Descarrega materials pedagògics i eines de codisseny desenvolupades durant la investigació.",
        "rec_btn_pdf": "Descarregar PDF (4.8 MB)",
        "rec_btn_zip": "Descarregar Plantilles (Zip, 12 MB)",
        "rec_btn_git": "Veure Repositori GitHub",
        "colab_title": "Amb la col·laboració i finançament de:",
        "footer_copy": "&copy; 2026 COPLITELE-IA. Projecte finançat pel Ministeri espanyol de Ciència i Innovació, desenvolupat pel GTE de la UIB i l'IRIE."
    },
    "en": {
        "menu_inicio": "Home",
        "menu_proyecto": "Project",
        "menu_impacto": "Impact & Dissemination",
        "hero_tag": "Active Investigation Project",
        "hero_main_title": "Redesigning the Future of Education with <span class=\"blue-highlight\">Codesign</span> & <span class=\"green-highlight\">Artificial Intelligence</span>",
        "hero_desc": "COPLITELE-IA is a research project transforming higher education by integrating Generative Artificial Intelligence (GAI) through an innovative pedagogical lens. Through educational co-design between faculty and students, we champion personalized learning using flexible itineraries, fostering connected virtual environments that ensure equity, inclusion, and educational quality.",
        "btn_conocer": "Explore the Project",
        "btn_publicaciones": "Publications",
        "news_title": "Latest news",
        "news_pretitle": "News",
        "stats_years": "Years of research",
        "stats_investigadores": "Researchers",
        "stats_publicaciones": "Publications",
        "stats_experiencias": "Experiencies",
        "progress_label": "Project Process",
        "submenu_desc": "Description",
        "submenu_obj": "Objectives",
        "submenu_miembros": "Members",
        "submenu_noticias": "Latest news",
        "submenu_transferencia": "Transfer",
        "submenu_publicaciones": "Scientific production",
        "submenu_recursos": "Resources",
        "obj_title": "Project Objectives",
        "obj_1_title": "Main Objective",
        "obj_1_desc": "Investigate the impact of AI in interlingual telecollaboration, analyzing how co-design improves language learning processes.",
        "obj_2_title": "Innovation",
        "obj_2_desc": "Develop adaptive personalization tools based on AI for technology-mediated language learning environments.",
        "obj_3_title": "International Scope",
        "obj_3_desc": "Collaboration with European and Latin American institutions to study computer-mediated communication in plurilingual contexts.",
        "obj_4_title": "Social Impact",
        "obj_4_desc": "Contribute to digital and educational inclusion through accessible and personalized technological solutions.",
        "project_pretitle": "Find out more",
        "project_title": "The Project",
        "tab_desc": "Description",
        "tab_equipo": "Team",
        "project_what_is": "What is COPLITELE-IA?",
        "project_what_is_p1": "<strong>COPLITELE-IA</strong> (Co-design, Personalization, and Technology in Telecollaboration for Language Learning with AI) is a research project funded by the Spanish Ministry of Science and Innovation (Ref. PID2021-127836NB-I00).",
        "project_what_is_p2": "The project examines how participatory co-design approaches, combined with artificial intelligence tools, can transform interlingual telecollaboration experiences, especially in foreign language learning contexts.",
        "project_what_is_p3": "Our approach integrates computer-mediated communication (CMC), computer-assisted language learning (CALL), and emerging AI technologies, including language models, sentiment analysis, and automated feedback systems.",
        "meta_ref": "Project Ref",
        "meta_duracion": "Duration",
        "meta_lider": "Lead Institution",
        "meta_financiacion": "Funding",
        "framework_header": "COPLITELE-IA",
        "framework_subtitle": "Research Framework",
        "trans_pretitle": "Impact and dissemination",
        "trans_title": "Transfer Activities",
        "trans_subtitle": "Activities designed to maximize the social and educational impact of our research, connecting academia with teaching practice and society.",
        "trans_btn_more": "See all activities",
        "pub_pretitle": "Scientific Library",
        "pub_title": "Scientific Production",
        "pub_subtitle": "Access project publications synchronized with our Zotero reference library.",
        "search_placeholder": "Search by title, author, or keyword...",
        "tab_all": "All",
        "tab_revistas": "Journals (Zotero)",
        "tab_libros": "Books (Zotero)",
        "tab_ponencias": "Presentations & Conferences",
        "rec_pretitle": "Open Repository",
        "rec_title": "Project Resources",
        "rec_subtitle": "Download pedagogical materials and co-design tools developed during the research.",
        "rec_btn_pdf": "Download PDF (4.8 MB)",
        "rec_btn_zip": "Download Templates (Zip, 12 MB)",
        "rec_btn_git": "View GitHub Repository",
        "colab_title": "With the collaboration and funding of:",
        "footer_copy": "&copy; 2026 COPLITELE-IA. Project funded by the Spanish Ministry of Science and Innovation, developed by the UIB GTE and IRIE."
    }
}

team_members = [
    {
        "id": "maria-gomez",
        "name": "Dra. Maria Gómez",
        "role": {
            "es": "Directora del Proyecto",
            "ca": "Directora del Projecte",
            "en": "Project Director"
        },
        "bio": {
            "es": "Doctora en Tecnología Educativa por la UIB. Especialista en metodologías participativas y codiseño de entornos inteligentes de aprendizaje. Coordina la vinculación del proyecto con las escuelas.",
            "ca": "Doctora en Tecnologia Educativa per la UIB. Especialista en metodologies participatives i codisseny d'entorns intel·ligents d'aprenentatge. Coordina la vinculació del projecte amb les escoles.",
            "en": "PhD in Educational Technology from UIB. Specialist in participatory methodologies and co-design of smart learning environments. She coordinates project partnerships with schools."
        },
        "email": "maria.gomez@uib.es",
        "orcid": "0000-0002-1823-9922"
    },
    {
        "id": "alberto-rodriguez",
        "name": "Dr. Alberto Rodríguez",
        "role": {
            "es": "Investigador Principal",
            "ca": "Investigador Principal",
            "en": "Lead Researcher"
        },
        "bio": {
            "es": "Doctor en Informática. Su campo de investigación es la Inteligencia Artificial aplicada a la personalización de contenidos y a la analítica de aprendizaje en educación superior.",
            "ca": "Doctor en Informàtica. El seu camp d'investigació és la Intel·ligència Artificial aplicada a la personalització de continguts i a l'analítica d'aprenentatge en educació superior.",
            "en": "PhD in Computer Science. His research field is Artificial Intelligence applied to content personalization and learning analytics in higher education."
        },
        "email": "alberto.rodriguez@uib.es",
        "orcid": "0000-0003-4902-8812"
    },
    {
        "id": "lucas-martinez",
        "name": "Dr. Lucas Martínez",
        "role": {
            "es": "Investigador en IA",
            "ca": "Investigador en IA",
            "en": "AI Researcher"
        },
        "bio": {
            "es": "Especialista en procesamiento del lenguaje natural y modelado de estudiante. Desarrolla los algoritmos predictivos adaptativos del proyecto.",
            "ca": "Especialista en processament del llenguatge natural i modelatge d'estudiant. Desenvolupa els algorismes predictius adaptatius del projecte.",
            "en": "Specialist in natural language processing and student modeling. He develops the project's predictive adaptive algorithms."
        },
        "email": "lucas.martinez@uib.es",
        "orcid": "0000-0001-9922-3844"
    },
    {
        "id": "sara-vidal",
        "name": "Sara Vidal",
        "role": {
            "es": "Diseñadora UX / Facilitadora",
            "ca": "Dissenyadora UX / Facilitadora",
            "en": "UX Designer / Facilitator"
        },
        "bio": {
            "es": "Máster en Diseño de Interacción. Dirige los talleres de codiseño con profesorado y alumnado, traduciendo requisitos pedagógicos en wireframes interactivos.",
            "ca": "Màster en Disseny d'Interacció. Dirigeix els tallers de codisseny amb professorat i alumnat, traduint requisits pedagògics en wireframes interactius.",
            "en": "Master in Interaction Design. Leads the co-design workshops with teachers and students, translating pedagogical requirements into interactive wireframes."
        },
        "email": "sara.vidal@uib.es",
        "orcid": "0000-0002-3811-1922"
    }
]

publications = [
    {
        "id": "pub-1",
        "type": "revista",
        "title": {
            "es": "Codiseño de entornos virtuales de aprendizaje personalizados mediante Inteligencia Artificial: Un enfoque cooperativo",
            "ca": "Codisseny d'entorns virtuals d'aprenentatge personalitzats mitjançant Intel·ligència Artificial: Un enfocament cooperatiu",
            "en": "Co-design of personalized virtual learning environments using Artificial Intelligence: A cooperative approach"
        },
        "citation": "Gómez, M., & Rodríguez, A. (2025). Revista de Educación y Tecnología, 14(2), 120-138.",
        "abstract": {
            "es": "Este artículo explora un marco metodológico para el codiseño de plataformas virtuales donde estudiantes y docentes participan activamente en la parametrización de algoritmos de inteligencia artificial para personalizar trayectorias de aprendizaje. Se detalla un estudio de caso en dos centros de secundaria y las percepciones de control de los usuarios frente al algoritmo.",
            "ca": "Aquest article explora un marc metodològic per al codisseny de plataformes virtuals on estudiants i docents participen activament en la parametrització d'algorismes d'intel·ligència artificial per personalitzar trajectòries d'aprenentatge. Es detalla un estudi de cas en dos centres de secundària i les percepcions de control dels usuaris enfront de l'algorisme.",
            "en": "This article explores a methodological framework for the co-design of virtual platforms where students and teachers actively participate in configuring artificial intelligence algorithms to personalize learning pathways. A case study in two secondary schools and users' perceptions of control over the algorithm are detailed."
        },
        "doi": "10.1016/j.edutec.2025.101230",
        "tags": ["Codiseño", "Inteligencia Artificial", "Educación"],
        "extraLabel": {
            "es": "Artículos",
            "ca": "Articles",
            "en": "Articles"
        },
        "zoteroKey": "GOM2025"
    },
    {
        "id": "pub-2",
        "type": "revista",
        "title": {
            "es": "La perspectiva de la comunidad educativa en el diseño de herramientas de IA: Desafíos prácticos de la co-creación",
            "ca": "La perspectiva de la comunitat educativa en el disseny d'eines d'IA: Desafiaments pràctics de la co-creació",
            "en": "The educational community's perspective on AI tool design: Practical challenges of co-creation"
        },
        "citation": "Martínez, L., Vidal, S., & Gómez, M. (2024). Pixel-Bit: Revista de Medios y Educación, 69, 45-78.",
        "abstract": {
            "es": "Estudio sobre los retos de comunicación y competencias tecnológicas que emergen al sentar en la misma mesa de codiseño a desarrolladores de software educativo e investigadores escolares. Se proponen dinámicas visuales para mitigar la asimetría técnica y empoderar a la comunidad educativa.",
            "ca": "Estudi sobre els reptes de comunicació i competències tecnològiques que emergeixen en seure a la mateixa taula de codisseny desenvolupadors de programari educatiu i investigadors escolars. Es proposen dinàmiques visuals per mitigar l'asimetria tècnica i empoderar la comunitat educativa.",
            "en": "Study on communication challenges and technical skills emerging when bringing educational software developers and school researchers together at the same co-design table. Visual dynamics are proposed to mitigate technical asymmetry and empower the educational community."
        },
        "doi": "10.12795/pixelbit.2024.10189",
        "tags": ["Co-creación", "Tecnología", "Usabilidad"],
        "extraLabel": {
            "es": "Artículos",
            "ca": "Articles",
            "en": "Articles"
        },
        "zoteroKey": "MAR2024"
    },
    {
        "id": "pub-3",
        "type": "libro",
        "title": {
            "es": "Tecnología Educativa y Personalización: Guía Práctica para el Codiseño de Aulas Inteligentes",
            "ca": "Tecnologia Educativa i Personalització: Guia Pràctica per al Codisseny d'Aules Intel·ligents",
            "en": "Educational Technology and Personalization: A Practical Guide for Co-designing Smart Classrooms"
        },
        "citation": "Gómez, M. (2024). Editorial UIB, Palma de Mallorca.",
        "abstract": {
            "es": "Un manual exhaustivo que provee marcos teóricos, plantillas de talleres de codiseño y guías éticas para la introducción de algoritmos adaptativos en el ámbito de la educación primaria y secundaria. Dirigido a formadores de profesorado y tecnólogos.",
            "ca": "Un manual exhaustiu que proveeix marcs teòrics, plantilles de tallers de codisseny i guies ètiques per a la introducció d'algorismes adaptatius en l'àmbit de l'educació primària i secundària. Adreçat a formadors de professorat i tecnòlegs.",
            "en": "A comprehensive manual providing theoretical frameworks, templates for co-design workshops, and ethical guidelines for implementing adaptive algorithms in primary and secondary education. Intended for teacher trainers and technologists."
        },
        "isbn": "978-84-8384-498-3",
        "tags": ["Manual", "Codiseño", "Aulas"],
        "extraLabel": {
            "es": "Publicaciones",
            "ca": "Publicacions",
            "en": "Publications"
        },
        "zoteroKey": "GOM2024b"
    },
    {
        "id": "pub-4",
        "type": "libro",
        "title": {
            "es": "Inteligencia Artificial y Educación: Nuevos horizontes para el codiseño docente",
            "ca": "Intel·ligència Artificial i Educació: Nous horitzons per al codisseny docent",
            "en": "Artificial Intelligence and Education: New horizons for teacher co-design"
        },
        "citation": "Rodríguez, A. (Ed.). (2025). Octaedro Editorial.",
        "abstract": {
            "es": "Una antología que reúne investigaciones iberoamericanas sobre el papel del docente como co-creador y supervisor de agentes inteligentes en el aula, discutiendo el diseño de cuadros de mando explicables y la soberanía del dato escolar.",
            "ca": "Una antologia que reuneix investigacions iberoamericanes sobre el paper del docent com a co-creador i supervisor d'agents intel·ligents a l'aula, discutint el disseny de quadres de comandament explicables i la sobirania de la dada escolar.",
            "en": "An anthology gathering Ibero-American research on the teacher's role as co-creator and supervisor of intelligent agents in the classroom, discussing the design of explainable dashboards and school data sovereignty."
        },
        "isbn": "978-84-19023-88-2",
        "tags": ["IA", "Docencia", "Innovación"],
        "extraLabel": {
            "es": "Publicaciones",
            "ca": "Publicacions",
            "en": "Publications"
        },
        "zoteroKey": "ROD2025b"
    },
    {
        "id": "pub-5",
        "type": "ponencia",
        "title": {
            "es": "Dynamic Interface Generation for Personalized Learning: A Co-design Case Study",
            "ca": "Dynamic Interface Generation for Personalized Learning: A Co-design Case Study",
            "en": "Dynamic Interface Generation for Personalized Learning: A Co-design Case Study"
        },
        "citation": "Presented at International Conference on Educational Technology (ICET), Paris, 2024.",
        "abstract": {
            "es": "Este artículo analiza la implementación técnica de interfaces configuradas dinámicamente a través de talleres de codiseño. Presentamos un marco de telemetría diseñado para equilibrar las recomendaciones automáticas de IA con los ajustes manuales del docente en tiempo real.",
            "ca": "Aquest article analitza la implementació tècnica d'interfícies configurades dinàmicament a través de tallers de codisseny. Presentem un marc de telemetria dissenyat per equilibrar les recomanacions automàtiques d'IA com els ajustaments manuals del docent en temps real.",
            "en": "This paper analyzes the technical implementation of interfaces configured dynamically through co-design workshops. We present a telemetry framework designed to balance agentic AI recommendations with manual teacher overrides in real-time."
        },
        "tags": ["UI", "Automation", "Agency"],
        "extraLabel": {
            "es": "Congresos",
            "ca": "Congressos",
            "en": "Conferences"
        },
        "zoteroKey": "ICET2024"
    },
    {
        "id": "pub-6",
        "type": "ponencia",
        "title": {
            "es": "El rol del codiseño en el desarrollo de asistentes virtuales inteligentes para secundaria",
            "ca": "El rol del codisseny en el desenvolupament d'assistents virtuals intel·ligents per a secundària",
            "en": "The role of co-design in the development of intelligent virtual assistants for secondary schools"
        },
        "citation": "Vidal, S. & Martínez, L. (2025). Ponencia en el Congreso Nacional de Investigación Educativa, Madrid.",
        "abstract": {
            "es": "Presentación de resultados del prototipado rápido de asistentes inteligentes en tres institutos de Mallorca, detallando la metodología de codiseño por fases (exploración, co-creación, evaluación) y la acogida de los tableros de control.",
            "ca": "Presentació de resultats del prototipat ràpid d'assistents intel·ligents en tres instituts de Mallorca, detallant la metodologia de codisseny per fases (exploració, co-creació, avaluació) i l'acollida dels quadres de comandament.",
            "en": "Presentation of results from rapid prototyping of intelligent assistants in three high schools in Mallorca, detailing the phased co-design methodology (exploration, co-creation, evaluation) and the acceptance of dashboard controls."
        },
        "tags": ["Asistentes", "Secundaria", "Prototipado"],
        "extraLabel": {
            "es": "Seminarios",
            "ca": "Seminaris",
            "en": "Seminars"
        },
        "zoteroKey": "VID2025"
    }
]

transfer_activities = [
    {
        "id": "act-1",
        "type": "taller",
        "tag": {"es": "Taller", "ca": "Taller", "en": "Workshop"},
        "title": {
            "es": "Taller de Codiseño con Docentes",
            "ca": "Taller de Codisseny amb Docents",
            "en": "Co-design Workshop with Teachers"
        },
        "desc": {
            "es": "Sesión participativa con profesorado de educación secundaria para co-crear estrategias de integración de IA en la enseñanza de lenguas.",
            "ca": "Sessió participativa amb professorat d'educació secundària per co-crear estratègies d'integració d'IA en l'ensenyament de llengües.",
            "en": "Participatory session with secondary education teachers to co-create strategies for integrating AI in language teaching."
        },
        "pills": ["Codiseño", "Formación Docente"],
        "date": "Marzo 2024",
        "location": "UIB, Palma de Mallorca",
        "loremIpsum": {
            "es": "El Taller de Codiseño con Docentes se concibió como un espacio participativo donde más de veinte profesores y profesoras de educación secundaria compartieron sus necesidades y expectativas sobre la integración de sistemas adaptativos de Inteligencia Artificial en el aula. A través de dinámicas de diseño centrado en el usuario, se definieron los requisitos para un panel de control que permita al docente supervisar las sugerencias de la IA en tiempo real. Este codiseño fomenta la soberanía educativa frente al algoritmo.\n\nDurante las sesiones de trabajo cooperativo, se elaboraron wireframes conceptuales de visualizaciones de datos y métricas adaptativas, determinando qué información sobre el procesamiento del lenguaje del alumnado resulta verdaderamente relevante para la toma de decisiones pedagógicas y cómo presentarla sin provocar sobrecarga cognitiva.",
            "ca": "El Taller de Codisseny amb Docents es va concebre com un espai participatiu on més de vint professors i professores d'educació secundària van compartir les seves necessitats i expectatives sobre la integració de sistemes adaptatius d'Intel·ligència Artificial a l'aula. A través de dinàmiques de disseny centrat en l'usuari, es van definir els requisits per a un panell de control que permeti al docent supervisar les suggeriments de la IA en temps real. Aquest codisseny fomenta la sobirania educativa enfront de l'algorisme.\n\nDurant les sessions de treball cooperatiu, es van elaborar wireframes conceptuals de visualitzacions de dades i mètriques adaptatives, determinant quina informació sobre el processament del llenguatge de l'alumnat resulta verdaderament rellevant per a la presa de decisions pedagògiques i com presentar-la sense provocar sobrecàrrega cognitiva.",
            "en": "The Co-design Workshop with Teachers was conceived as a participatory space where more than twenty secondary education teachers shared their needs and expectations regarding the integration of adaptive Artificial Intelligence systems in the classroom. Through user-centered design dynamics, requirements were defined for a dashboard that allows teachers to supervise AI suggestions in real time. This co-design promotes educational sovereignty in the face of the algorithm.\n\nDuring the cooperative work sessions, conceptual wireframes of data visualizations and adaptive metrics were developed, determining what information on student language processing is truly relevant for pedagogical decision-making and how to present it without causing cognitive overload."
        }
    },
    {
        "id": "act-2",
        "type": "seminario",
        "tag": {"es": "Seminario", "ca": "Seminari", "en": "Seminar"},
        "title": {
            "es": "Seminario Internacional: IA y CALL",
            "ca": "Seminari Internacional: IA i CALL",
            "en": "International Seminar: AI and CALL"
        },
        "desc": {
            "es": "Seminario online internacional con expertos de Europa y América Latina sobre el estado del arte en IA aplicada al aprendizaje de lenguas.",
            "ca": "Seminari online internacional amb experts d'Europa i Amèrica Llatina sobre l'estat de l'art en IA aplicada a l'aprenentatge de llengües.",
            "en": "International online seminar with experts from Europe and Latin America on the state of the art in AI applied to language learning."
        },
        "pills": ["IA", "CALL", "Internacional"],
        "date": "Mayo 2024",
        "location": "Online (Zoom)",
        "loremIpsum": {
            "es": "El Seminario Internacional: IA y CALL congregó a investigadores de primer nivel procedentes de universidades europeas y americanas. Las sesiones se centraron en debatir el estado del arte de las tecnologías de procesamiento del lenguaje natural (NLP) aplicadas a la adquisición de segundas lenguas (CALL), prestando especial atención a la retroalimentación automática y al análisis de interacciones en contextos de telecolaboración virtual.\n\nLos ponentes discutieron sobre los límites éticos de los modelos de lenguaje comerciales en entornos escolares y presentaron propuestas de implementación local que salvaguardan la privacidad del estudiante. El debate evidenció la necesidad de crear frameworks híbridos que unan la intuición pedagógica humana con la precisión de los análisis basados en deep learning.",
            "ca": "El Seminari Internacional: IA i CALL va congregar investigadors de primer nivell procedents d'universitats europees i americanes. Les sessions es van centrar en debatre l'estat de l'art de les tecnologies de processament del llenguatge natural (NLP) aplicades a l'adquisició de segones llengües (CALL), prestant especial atenció a la retroalimentació automàtica i a l'anàlisi d'interaccions en contextos de telecol·laboració virtual.\n\nEls ponents van discutir sobre els límits ètics dels models de llenguatge comercials en entorns escolars i van presentar propostes d'implementació local que salvaguarden la privacitat de l'estudiant. El debat va evidenciar la necessitat de crear frameworks híbrids que uneixin la intuïció pedagògica humana amb la precisió de les anàlisis basades en deep learning.",
            "en": "The International Seminar: AI and CALL gathered top-tier researchers from European and American universities. The sessions focused on debating the state of the art of natural language processing (NLP) technologies applied to computer-assisted language learning (CALL), paying special attention to automated feedback and interaction analysis in virtual telecollaboration contexts.\n\nSpeakers discussed the ethical limits of commercial language models in school environments and presented local implementation proposals that safeguard student privacy. The debate highlighted the need to create hybrid frameworks that combine human pedagogical intuition with the precision of deep learning-based analytics."
        }
    },
    {
        "id": "act-3",
        "type": "demo",
        "tag": {"es": "Demo", "ca": "Demo", "en": "Demo"},
        "title": {
            "es": "Demostración de Prototipo",
            "ca": "Demostració de Prototip",
            "en": "Prototype Demonstration"
        },
        "desc": {
            "es": "Presentación del prototipo de plataforma de telecolaboración con módulo de IA integrado. Sesión abierta a la comunidad universitaria.",
            "ca": "Presentació del prototip de plataforma de telecolaboració amb mòdul d'IA integrat. Sessió oberta a la comunitat universitària.",
            "en": "Presentation of the telecollaboration platform prototype with integrated AI module. Session open to the university community."
        },
        "pills": ["Tecnología", "Demostración"],
        "date": "Octubre 2024",
        "location": "UIB, Palma de Mallorca",
        "loremIpsum": {
            "es": "La Demostración de Prototipo del sistema interactivo COPLITELE-IA sirvió para validar con la comunidad investigadora y educativa la primera versión funcional del entorno adaptativo de telecolaboración. El prototipo permite parametrizar de forma ágil la retroalimentación predictiva del sistema, proporcionando una interfaz limpia y accesible que asegura que la mediación tecnológica sea explicable y transparente.\n\nLos asistentes pudieron probar de primera mano cómo los alumnos interactúan en las salas de telecolaboración mientras un motor de análisis lingüístico proporciona sugerencias dinámicas adaptadas al nivel de competencia de cada participante. Las respuestas recolectadas servirán para pulir los tiempos de respuesta y la precisión de la interfaz en futuras iteraciones del proyecto.",
            "ca": "La Demostració de Prototip del sistema interactiu COPLITELE-IA va servir per validar amb la comunitat investigadora i educativa la primera versió funcional de l'entorn adaptatiu de telecol·laboració. El prototip permet parametritzar de manera àgil la retroalimentació predictiva del sistema, proporcionant una interfície neta i accessible que assegura que la mediació tecnològica sigui explicable i transparent.\n\nEls assistents van poder provar de primera mà com els alumnes interactuen a les sales de telecol·laboració mentre un motor d'anàlisi lingüística proporciona suggeriments dinàmics adaptats al nivell de competència de cada participant. Les respostes recollides serviran per polir els temps de resposta i la precisió de la interfície en futures iteracions del projecte.",
            "en": "The Prototype Demonstration of the COPLITELE-IA interactive system served to validate the first functional version of the adaptive telecollaboration environment with the research and educational community. The prototype allows for agile configuration of the system's predictive feedback, providing a clean and accessible interface that ensures the technological mediation is explainable and transparent.\n\nAttendees experienced firsthand how students interact in the telecollaboration rooms while a linguistic analysis engine provides dynamic suggestions adapted to each participant's proficiency level. The feedback collected will be used to refine response times and interface precision in future project iterations."
        }
    },
    {
        "id": "act-4",
        "type": "formacion",
        "tag": {"es": "Formación", "ca": "Formació", "en": "Training"},
        "title": {
            "es": "Curso de Formación: NLP para Educadores",
            "ca": "Curs de Formació: NLP per a Educadors",
            "en": "Training Course: NLP for Educators"
        },
        "desc": {
            "es": "Programa de formación de 8 horas dirigido a docentes de lenguas sobre procesamiento del lenguaje natural y herramientas de feedback automatizado.",
            "ca": "Programa de formació de 8 hores adreçat a docents de llengües sobre processament del llenguatge natural i eines de feedback automatitzat.",
            "en": "8-hour training program for language teachers on natural language processing and automated feedback tools."
        },
        "pills": ["NLP", "Formación", "Blended"],
        "date": "Enero 2025",
        "location": "Online y presencial (UIB)",
        "loremIpsum": {
            "es": "El Curso de Formación en NLP para Educadores proporcionó al profesorado participante herramientas conceptuales y prácticas para comprender cómo la IA procesa el lenguaje de los estudiantes. Durante el curso, se diseñaron actividades de escritura que incorporaban sistemas automáticos de retroalimentación constructiva, analizando el impacto de este andamiaje algorítmico en la motivación del alumnado.\n\nSe profundizó en el funcionamiento básico de los transformadores y los modelos de lenguaje, explicando cómo traducir la probabilidad de palabras en andamios de escritura significativos. Los docentes participantes diseñaron guías didácticas que integran estas herramientas como asistentes de escritura, promoviendo el pensamiento crítico frente al texto generado automáticamente.",
            "ca": "El Curs de Formació en NLP per a Educadors va proporcionar al professorat participant eines conceptuals i pràctiques per comprendre com la IA processa el llenguatge dels estudiants. Durant el curs, es van dissenyar activitats d'escriptura que incorporaven sistemes automàtics de retroalimentació constructiva, analitzant l'impacte d'aquesta bastida algorísmica en la motivació de l'alumnat.\n\nEs va aprofundir en el funcionament bàsic dels transformadors i els models de llenguatge, explicant com traduir la probabilitat de paraules en bastides d'escriptura significatives. Els docents participants van dissenyar guies didàctiques que integren aquestes eines com a assistents d'escriptura, promovent el pensament crític enfront del text generat automàticament.",
            "en": "The NLP for Educators Training Course provided participating teachers with conceptual and practical tools to understand how AI processes student language. During the course, writing activities incorporating automatic systems of constructive feedback were designed, analyzing the impact of this algorithmic scaffolding on student motivation.\n\nThe basic operation of transformers and language models was covered, explaining how to translate word probabilities into meaningful writing scaffolds. Participating teachers designed educational guides that integrate these tools as writing assistants, promoting critical thinking toward automatically generated text."
        }
    },
    {
        "id": "act-5",
        "type": "taller",
        "tag": {"es": "Taller", "ca": "Taller", "en": "Workshop"},
        "title": {
            "es": "Workshop: Diseño de Tareas Telecolaborativas",
            "ca": "Workshop: Disseny de Tasques Telecolaboratives",
            "en": "Workshop: Designing Telecollaborative Tasks"
        },
        "desc": {
            "es": "Taller práctico para el diseño colaborativo de tareas de intercambio lingüístico entre estudiantes de diferentes países con apoyo de IA.",
            "ca": "Taller pràctic per al disseny col·laboratiu de tasques d'intercanvi lingüístic entre estudiants de diferents països amb suport d'IA.",
            "en": "Practical workshop for the collaborative design of language exchange tasks between students from different countries supported by AI."
        },
        "pills": ["Diseño de Tareas", "Telecolaboración"],
        "date": "Marzo 2025",
        "location": "UAH, Alcalá de Henares",
        "loremIpsum": {
            "es": "El Workshop de Diseño de Tareas Telecolaborativas, celebrado en Alcalá de Henares, sirvió como un foro de co-creación de itinerarios didácticos virtuales. Los investigadores y docentes desarrollaron tareas de comunicación interactiva apoyadas por IA, adaptadas a diversos niveles de competencia lingüística, priorizando la equidad de las dinámicas en contextos plurilingües.\n\nLas dinámicas permitieron unificar criterios sobre cómo estructurar los turnos de intervención de los alumnos, y cómo la plataforma inteligente puede sugerir disparadores de conversación (prompts) en caso de bloqueos comunicativos, asegurando que las parejas mantengan un flujo natural de telecolaboración.",
            "ca": "El Workshop de Disseny de Tasques Telecol·laboratives, celebrat a Alcalá de Henares, va servir com un fòrum de co-creació d'itineraris didàctics virtuals. Els investigadors i docents van desenvolupar tasques de comunicació interactiva amb suport d'IA, adaptades a diversos nivells de competència lingüística, prioritzant l'equitat de les dinàmiques en contextos plurilingües.\n\nLes dinàmiques van permetre unificar criteris sobre com estructurar els torns d'intervenció dels alumnes, i com la plataforma intel·ligent pot suggerir disparadors de conversa (prompts) en cas de bloquejos comunicatius, assegurant que les parelles mantinguin un flux natural de telecol·laboració.",
            "en": "The Designing Telecollaborative Tasks Workshop, held in Alcalá de Henares, served as a forum for co-creating virtual educational itineraries. Researchers and teachers developed interactive communication tasks supported by AI, adapted to various language proficiency levels, prioritizing equity in plurilingual contexts.\n\nThe dynamics allowed for unifying criteria on how to structure student turn-taking, and how the smart platform can suggest conversation starters (prompts) in case of communication blocks, ensuring that partners maintain a natural flow of telecollaboration."
        }
    },
    {
        "id": "act-6",
        "type": "difusion",
        "tag": {"es": "Difusión", "ca": "Difusió", "en": "Dissemination"},
        "title": {
            "es": "Jornada de Difusión de Resultados",
            "ca": "Jornada de Difusió de Resultats",
            "en": "Results Dissemination Day"
        },
        "desc": {
            "es": "Presentación pública de los principales hallazgos del proyecto, con mesas redondas y presentación de materiales didácticos elaborados.",
            "ca": "Presentació pública de les principals troballes del projecte, amb taules rodones i presentació de materials didàctics elaborats.",
            "en": "Public presentation of the project's main findings, with roundtables and presentation of developed educational materials."
        },
        "pills": ["Resultados", "Difusión"],
        "date": "Junio 2025",
        "location": "UIB, Palma de Mallorca",
        "loremIpsum": {
            "es": "La Jornada de Difusión de Resultados marcó el cierre público del proyecto, reuniendo a autoridades académicas, docentes y público general. Se presentaron los principales hallazgos de las pruebas de campo, demostrando el impacto positivo de la telecolaboración asistida por IA co-diseñada, y se distribuyeron manuales didácticos y códigos de buenas prácticas éticas.\n\nLas ponencias resumieron las métricas cuantitativas y cualitativas registradas en los institutos, destacando una mejora estadísticamente significativa en la fluidez de escritura de los estudiantes y una reducción de la ansiedad lingüística gracias al andamiaje personalizado del sistema adaptativo.",
            "ca": "La Jornada de Difusió de Resultats va marcar el tancament públic del projecte, reunint autoritats acadèmiques, docents i públic general. Es van presentar les principals troballes de les proves de camp, demostrant l'impacte positiu de la telecol·laboració assistida per IA co-dissenyada, i es van distribuir manuals didàctics i codis de bones pràctiques ètiques.\n\nLes ponències van resumir les mètriques quantitatives i qualitatives registrades als instituts, destacant una millora estadísticament significativa en la fluïdesa d'escriptura dels estudiants i una reducció de l'ansietat lingüística gràcies a la bastida personalitzada del sistema adaptatiu.",
            "en": "The Results Dissemination Day marked the public closing of the project, bringing together academic authorities, teachers, and the general public. The main findings from the field trials were presented, demonstrating the positive impact of co-designed AI-assisted telecollaboration, and educational manuals and codes of ethical good practices were distributed.\n\nPresentations summarized both quantitative and qualitative metrics recorded in high schools, highlighting a statistically significant improvement in students' writing fluency and a reduction in language anxiety due to the personalized scaffolding of the adaptive system."
        }
    }
]

# Hardcoded resources in index.html (Spanish-only text descriptions)
resources_hardcoded = [
    {
        "title": "Guía de Codiseño en Tecnología Educativa",
        "category": "Manual Metodológico",
        "desc": "Un compendio de dinámicas, fichas imprimibles y directrices prácticas para planificar y llevar a cabo talleres de codiseño tecnológico con profesorado y estudiantes.",
        "btn_key": "rec_btn_pdf"
    },
    {
        "title": "Fichas de Explicabilidad Algorítmica",
        "category": "Plantillas Didácticas",
        "desc": "Materiales didácticos de apoyo para trabajar la soberanía y comprensión del dato escolar con alumnos de secundaria y bachillerato en sesiones presenciales.",
        "btn_key": "rec_btn_zip"
    },
    {
        "title": "Módulo de Personalización Soberana",
        "category": "Código Abierto",
        "desc": "Accede al repositorio público del prototipo funcional que integra las librerías de personalización algorítmica parametrizable con interfaces web adaptables.",
        "btn_key": "rec_btn_git"
    }
]


def generate_markdown():
    """Generates a clean Markdown file with translation comparison"""
    md = []
    md.append("# COPLITELE-IA - Secciones de la Web para Verificación de Clientes")
    md.append("Este documento recopila la estructura completa y el contenido textual de cada sección de la página web de **COPLITELE-IA** en sus tres idiomas (Español, Catalán e Inglés). Su finalidad es facilitar la lectura, revisión y validación de la información con los clientes.")
    md.append("")
    md.append("---")
    md.append("")
    md.append("## ÍNDICE")
    md.append("1. [Identidad Visual y Conceptos de Marca](#1-identidad-visual-y-conceptos-de-marca)")
    md.append("2. [Textos de Navegación y Generales](#2-textos-de-navegacion-y-generales)")
    md.append("3. [Vista: Inicio (Home)](#3-vista-inicio-home)")
    md.append("4. [Vista: Proyecto](#4-vista-proyecto)")
    md.append("5. [Vista: Impacto y Difusión](#5-vista-impacto-y-difusion)")
    md.append("")
    md.append("---")
    md.append("")
    md.append("## 1. IDENTIDAD VISUAL Y CONCEPTOS DE MARCA")
    md.append("El logotipo dinámico de la web incorpora tres ejes de desarrollo en su animación, cada uno asociado a un color y tres palabras conceptuales en inglés, además de un lema de soporte:")
    md.append("")
    md.append("- **Eje Azul (Codiseño)**:")
    md.append("  - Palabras en logo: `Codesign`, `Learning`, `Itineraries`")
    md.append("  - Lema en cabecera: **Codiseño de aprendizaje flexible**")
    md.append("- **Eje Celeste (Personalización)**:")
    md.append("  - Palabras en logo: `Enhance`, `Personalized`, `Environments`")
    md.append("  - Lema en cabecera: **Itinerarios personalizados y agénticos**")
    md.append("- **Eje Verde (Tecnología)**:")
    md.append("  - Palabras en logo: `Technology`, `Inteligencia`, `Artificial`")
    md.append("  - Lema en cabecera: **Ambientes enriquecidos con Tecnología**")
    md.append("")
    md.append("---")
    md.append("")
    md.append("## 2. TEXTOS DE NAVEGACIÓN Y GENERALES")
    md.append("Estos textos corresponden al menú superior, botones de control de idioma, barra de búsqueda y pie de página.")
    md.append("")
    
    navigation_keys = [
        ("menu_inicio", "Opción Inicio del Menú"),
        ("menu_proyecto", "Opción Proyecto del Menú"),
        ("menu_impacto", "Opción Impacto y Difusión"),
        ("search_placeholder", "Placeholders de búsqueda"),
        ("colab_title", "Título Colaboradores"),
        ("footer_copy", "Copyright y Pie de página")
    ]
    
    for key, label in navigation_keys:
        md.append(f"### ID: `{key}` ({label})")
        md.append(f"- **Español (ES)**: {translations['es'][key]}")
        md.append(f"- **Catalán (CA)**: {translations['ca'][key]}")
        md.append(f"- **Inglés (EN)**: {translations['en'][key]}")
        md.append("")
        
    md.append("---")
    md.append("")
    md.append("## 3. VISTA: INICIO (HOME)")
    md.append("Se compone del mensaje de bienvenida principal (Hero), estadísticas y llamadas a la acción.")
    md.append("")
    
    home_keys = [
        ("hero_tag", "Etiqueta superior del Hero"),
        ("hero_main_title", "Título Principal del Hero"),
        ("hero_desc", "Descripción detallada del Hero"),
        ("btn_conocer", "Texto del Botón Principal"),
        ("progress_label", "Etiqueta de la barra de progreso (Valor fijo: 75%)"),
        ("stats_years", "Estadística 1: Años de Investigación"),
        ("stats_investigadores", "Estadística 2: Investigadores"),
        ("stats_publicaciones", "Estadística 3: Publicaciones"),
        ("stats_experiencias", "Estadística 4: Experiencias en centros"),
        ("news_title", "Cabecera de noticias en Inicio")
    ]
    
    for key, label in home_keys:
        md.append(f"### ID: `{key}` ({label})")
        md.append(f"- **Español (ES)**: {translations['es'][key].replace('<strong>', '**').replace('</strong>', '**')}")
        md.append(f"- **Catalán (CA)**: {translations['ca'][key].replace('<strong>', '**').replace('</strong>', '**')}")
        md.append(f"- **Inglés (EN)**: {translations['en'][key].replace('<strong>', '**').replace('</strong>', '**')}")
        md.append("")
        
    md.append("---")
    md.append("")
    md.append("## 4. VISTA: PROYECTO")
    md.append("Esta vista se organiza bajo un menú interno con tres pestañas: Descripción, Objetivos y Miembros.")
    md.append("")
    md.append("### Submenú Interno (Proyecto)")
    md.append(f"- **Descripción (`submenu_desc`)**: ES: \"{translations['es']['submenu_desc']}\" | CA: \"{translations['ca']['submenu_desc']}\" | EN: \"{translations['en']['submenu_desc']}\"")
    md.append(f"- **Objetivos (`submenu_obj`)**: ES: \"{translations['es']['submenu_obj']}\" | CA: \"{translations['ca']['submenu_obj']}\" | EN: \"{translations['en']['submenu_obj']}\"")
    md.append(f"- **Miembros (`submenu_miembros`)**: ES: \"{translations['es']['submenu_miembros']}\" | CA: \"{translations['ca']['submenu_miembros']}\" | EN: \"{translations['en']['submenu_miembros']}\"")
    md.append("")
    
    md.append("### 4.1. Sección: Descripción del Proyecto")
    
    desc_keys = [
        ("project_pretitle", "Pretítulo de sección"),
        ("project_title", "Título de sección"),
        ("project_what_is", "Pregunta de cabecera"),
        ("project_what_is_p1", "Primer párrafo descriptivo"),
        ("project_what_is_p2", "Segundo párrafo descriptivo"),
        ("project_what_is_p3", "Tercer párrafo descriptivo")
    ]
    
    for key, label in desc_keys:
        md.append(f"#### ID: `{key}` ({label})")
        md.append(f"- **Español (ES)**: {translations['es'][key].replace('<strong>', '**').replace('</strong>', '**')}")
        md.append(f"- **Catalán (CA)**: {translations['ca'][key].replace('<strong>', '**').replace('</strong>', '**')}")
        md.append(f"- **Inglés (EN)**: {translations['en'][key].replace('<strong>', '**').replace('</strong>', '**')}")
        md.append("")
        
    md.append("#### Ficha Técnica de Descripción (Metadatos)")
    meta_keys = [
        ("meta_ref", "Ref. Proyecto", "PID2021-127836NB-I00"),
        ("meta_duracion", "Duración", "2022-2025"),
        ("meta_lider", "Institución Líder", "UIB (Universitat de les Illes Balears)"),
        ("meta_financiacion", "Financiación", "MCIN / AEI")
    ]
    for key, label, val in meta_keys:
        md.append(f"- **{label} (`{key}`)**:")
        md.append(f"  - Etiqueta: ES: \"{translations['es'][key]}\" | CA: \"{translations['ca'][key]}\" | EN: \"{translations['en'][key]}\"")
        md.append(f"  - Valor: `{val}`")
    md.append("")
    
    md.append("#### Tarjeta del Framework Metodológico (Visual)")
    md.append(f"- **Título (`framework_header`)**: ES/CA/EN: \"{translations['es']['framework_header']}\"")
    md.append(f"- **Subtítulo (`framework_subtitle`)**: ES: \"{translations['es']['framework_subtitle']}\" | CA: \"{translations['ca']['framework_subtitle']}\" | EN: \"{translations['en']['framework_subtitle']}\"")
    md.append("- **Etiquetas fijas (Temáticas)**: `Codiseño` | `Telecolaboración` | `IA & NLP` | `CALL` | `Personalización` | `CMC`")
    md.append("")
    
    md.append("### 4.2. Sección: Objetivos del Proyecto")
    md.append(f"**Título de la sección (`obj_title`)**: ES: \"{translations['es']['obj_title']}\" | CA: \"{translations['ca']['obj_title']}\" | EN: \"{translations['en']['obj_title']}\"")
    md.append("")
    
    obj_items = [
        ("obj_1_title", "obj_1_desc", "Objetivo 1 (Principal)"),
        ("obj_2_title", "obj_2_desc", "Objetivo 2 (Innovación)"),
        ("obj_3_title", "obj_3_desc", "Objetivo 3 (Alcance Internacional)"),
        ("obj_4_title", "obj_4_desc", "Objetivo 4 (Impacto Social)")
    ]
    
    for t_key, d_key, label in obj_items:
        md.append(f"#### {label}")
        md.append(f"- **Título (`{t_key}`)**:")
        md.append(f"  - ES: \"{translations['es'][t_key]}\" | CA: \"{translations['ca'][t_key]}\" | EN: \"{translations['en'][t_key]}\"")
        md.append(f"- **Descripción (`{d_key}`)**:")
        md.append(f"  - ES: \"{translations['es'][d_key]}\"")
        md.append(f"  - CA: \"{translations['ca'][d_key]}\"")
        md.append(f"  - EN: \"{translations['en'][d_key]}\"")
        md.append("")
        
    md.append("### 4.3. Sección: Equipo de Investigadores (`equipo`)")
    md.append(f"**Título de la sección (`tab_equipo`)**: ES: \"{translations['es']['tab_equipo']}\" | CA: \"{translations['ca']['tab_equipo']}\" | EN: \"{translations['en']['tab_equipo']}\"")
    md.append("")
    
    for member in team_members:
        md.append(f"#### Miembro: {member['name']} (ID: `{member['id']}`)")
        md.append(f"- **Rol**:")
        md.append(f"  - ES: \"{member['role']['es']}\" | CA: \"{member['role']['ca']}\" | EN: \"{member['role']['en']}\"")
        md.append(f"- **Biografía**:")
        md.append(f"  - **ES**: {member['bio']['es']}")
        md.append(f"  - **CA**: {member['bio']['ca']}")
        md.append(f"  - **EN**: {member['bio']['en']}")
        md.append(f"- **Contacto**: Email: {member['email']} | ORCID: {member['orcid']}")
        md.append("")
        
    md.append("---")
    md.append("")
    md.append("## 5. VISTA: IMPACTO Y DIFUSIÓN")
    md.append("Contiene las noticias detalladas, actividades de transferencia docente, artículos de Zotero y recursos descargables.")
    md.append("")
    md.append("### Submenú Interno (Impacto)")
    md.append(f"- **Noticias (`submenu_noticias`)**: ES: \"{translations['es']['submenu_noticias']}\" | CA: \"{translations['ca']['submenu_noticias']}\" | EN: \"{translations['en']['submenu_noticias']}\"")
    md.append(f"- **Transferencia (`submenu_transferencia`)**: ES: \"{translations['es']['submenu_transferencia']}\" | CA: \"{translations['ca']['submenu_transferencia']}\" | EN: \"{translations['en']['submenu_transferencia']}\"")
    md.append(f"- **Producción Científica (`submenu_publicaciones`)**: ES: \"{translations['es']['submenu_publicaciones']}\" | CA: \"{translations['ca']['submenu_publicaciones']}\" | EN: \"{translations['en']['submenu_publicaciones']}\"")
    md.append(f"- **Recursos (`submenu_recursos`)**: ES: \"{translations['es']['submenu_recursos']}\" | CA: \"{translations['ca']['submenu_recursos']}\" | EN: \"{translations['en']['submenu_recursos']}\"")
    md.append("")
    
    md.append("### 5.1. Actividades de Transferencia (Detalle Completo)")
    md.append(f"**Título de la sección**: ES: \"{translations['es']['trans_title']}\"")
    md.append(f"**Subtítulo de la sección (`trans_subtitle`)**: ES: \"{translations['es']['trans_subtitle']}\" | CA: \"{translations['ca']['trans_subtitle']}\" | EN: \"{translations['en']['trans_subtitle']}\"")
    md.append("")
    
    for act in transfer_activities:
        md.append(f"#### Actividad: {act['title']['es']} (ID: `{act['id']}`)")
        md.append(f"- **Tipo / Categoría**: ES: \"{act['tag']['es']}\" | CA: \"{act['tag']['ca']}\" | EN: \"{act['tag']['en']}\"")
        md.append(f"- **Título Multilingüe**:")
        md.append(f"  - ES: \"{act['title']['es']}\"")
        md.append(f"  - CA: \"{act['title']['ca']}\"")
        md.append(f"  - EN: \"{act['title']['en']}\"")
        md.append(f"- **Descripción Breve**:")
        md.append(f"  - **ES**: {act['desc']['es']}")
        md.append(f"  - **CA**: {act['desc']['ca']}")
        md.append(f"  - **EN**: {act['desc']['en']}")
        md.append(f"- **Datos**: Fecha: {act['date']} | Ubicación: {act['location']} | Etiquetas: {', '.join(act['pills'])}")
        md.append("- **Texto de Detalle Completo (Desplegado en Modal / Ficha Individual)**:")
        md.append(f"  - **ES**: {act['loremIpsum']['es']}")
        md.append(f"  - **CA**: {act['loremIpsum']['ca']}")
        md.append(f"  - **EN**: {act['loremIpsum']['en']}")
        md.append("")
        
    md.append("### 5.2. Producción Científica (Publicaciones vinculadas a Zotero)")
    md.append(f"**Título de la sección (`pub_title`)**: ES: \"{translations['es']['pub_title']}\" | CA: \"{translations['ca']['pub_title']}\" | EN: \"{translations['en']['pub_title']}\"")
    md.append(f"**Subtítulo de la sección (`pub_subtitle`)**: ES: \"{translations['es']['pub_subtitle']}\" | CA: \"{translations['ca']['pub_subtitle']}\" | EN: \"{translations['en']['pub_subtitle']}\"")
    md.append("")
    
    for pub in publications:
        md.append(f"#### Publicación: {pub['title']['es']} (ID: `{pub['id']}`)")
        md.append(f"- **Tipo**: ES: \"{pub['extraLabel']['es']}\" | CA: \"{pub['extraLabel']['ca']}\" | EN: \"{pub['extraLabel']['en']}\"")
        md.append(f"- **Cita Bibliográfica**: {pub['citation']}")
        md.append(f"- **Título Multilingüe**:")
        md.append(f"  - ES: \"{pub['title']['es']}\"")
        md.append(f"  - CA: \"{pub['title']['ca']}\"")
        md.append(f"  - EN: \"{pub['title']['en']}\"")
        md.append(f"- **Resumen (Abstract)**:")
        md.append(f"  - **ES**: {pub['abstract']['es']}")
        md.append(f"  - **CA**: {pub['abstract']['ca']}")
        md.append(f"  - **EN**: {pub['abstract']['en']}")
        extra_info = []
        if 'doi' in pub: extra_info.append(f"DOI: {pub['doi']}")
        if 'isbn' in pub: extra_info.append(f"ISBN: {pub['isbn']}")
        extra_info.append(f"Zotero Key: {pub['zoteroKey']}")
        md.append(f"- **Metadatos**: {' | '.join(extra_info)}")
        md.append(f"- **Palabras clave**: {', '.join(pub['tags'])}")
        md.append("")
        
    md.append("### 5.3. Recursos del Proyecto (Descargables)")
    md.append(f"**Título de la sección (`rec_title`)**: ES: \"{translations['es']['rec_title']}\" | CA: \"{translations['ca']['rec_title']}\" | EN: \"{translations['en']['rec_title']}\"")
    md.append(f"**Subtítulo de la sección (`rec_subtitle`)**: ES: \"{translations['es']['rec_subtitle']}\" | CA: \"{translations['ca']['rec_subtitle']}\" | EN: \"{translations['en']['rec_subtitle']}\"")
    md.append("")
    
    for res in resources_hardcoded:
        md.append(f"#### Recurso: {res['title']}")
        md.append(f"- **Categoría**: {res['category']}")
        md.append(f"- **Descripción (ES - General)**: {res['desc']}")
        btn_k = res['btn_key']
        md.append(f"- **Botón de Descarga (`{btn_k}`)**: ES: \"{translations['es'][btn_k]}\" | CA: \"{translations['ca'][btn_k]}\" | EN: \"{translations['en'][btn_k]}\"")
        md.append("")
        
    return "\n".join(md)


def generate_html():
  """Generates a styled HTML file with navigation, grids, and nice margins for copy-pasting to Google Docs"""
  html = []
  html.append("""<!DOCTYPE html>
<html lang="es">
<head>
  <meta charset="UTF-8">
  <title>COPLITELE-IA - Secciones de la Web para Verificación de Clientes</title>
  <style>
    body {
      font-family: 'Georgia', 'Times New Roman', serif;
      line-height: 1.6;
      color: #333333;
      margin: 40px auto;
      max-width: 900px;
      padding: 0 20px;
    }
    h1 {
      font-family: 'Helvetica Neue', Arial, sans-serif;
      color: #1D5BFE;
      border-bottom: 2px solid #1D5BFE;
      padding-bottom: 10px;
      margin-top: 40px;
    }
    h2 {
      font-family: 'Helvetica Neue', Arial, sans-serif;
      color: #0f766e;
      border-bottom: 1px solid #e2e8f0;
      padding-bottom: 5px;
      margin-top: 35px;
    }
    h3 {
      font-family: 'Helvetica Neue', Arial, sans-serif;
      color: #1e293b;
      margin-top: 25px;
      margin-bottom: 10px;
    }
    h4 {
      font-family: 'Helvetica Neue', Arial, sans-serif;
      color: #475569;
      margin-top: 20px;
      margin-bottom: 5px;
      font-size: 1.1em;
    }
    .introduction {
      background-color: #f8fafc;
      border-left: 4px solid #1D5BFE;
      padding: 15px 20px;
      margin-bottom: 30px;
      font-style: italic;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      margin: 20px 0;
      font-size: 14px;
    }
    th, td {
      border: 1px solid #cbd5e1;
      padding: 10px;
      text-align: left;
    }
    th {
      background-color: #f1f5f9;
      font-weight: bold;
    }
    .lang-section {
      margin-left: 20px;
      margin-bottom: 15px;
      border-left: 3px solid #e2e8f0;
      padding-left: 15px;
    }
    .lang-tag {
      font-weight: bold;
      color: #2563eb;
      text-transform: uppercase;
      font-size: 0.9em;
      margin-top: 5px;
    }
    .meta-box {
      background-color: #f8fafc;
      padding: 10px 15px;
      border: 1px solid #e2e8f0;
      border-radius: 4px;
      margin: 10px 0;
      font-size: 0.95em;
    }
    .badge {
      display: inline-block;
      padding: 2px 8px;
      background-color: #e0f2fe;
      color: #0369a1;
      border-radius: 4px;
      font-size: 0.85em;
      font-weight: bold;
    }
    hr {
      border: 0;
      height: 1px;
      background: #cbd5e1;
      margin: 40px 0;
    }
  </style>
</head>
<body>

  <h1>COPLITELE-IA - Secciones de la Web para Verificación de Clientes</h1>
  <div class="introduction">
    Este documento sirve para verificar toda la información y textos de la web de COPLITELE-IA en español, catalán e inglés de cara a la validación con los clientes. El texto conservará su estructura clara al ser copiado y pegado en Google Docs o al importarse directamente.
  </div>

  <h2>1. Identidad Visual y Conceptos de Marca</h2>
  <p>El logotipo dinámico de la web incorpora tres ejes de desarrollo en su animación, cada uno asociado a un color y tres palabras conceptuales en inglés, además de un lema de soporte:</p>
  <ul>
    <li><strong>Eje Azul (Codiseño)</strong>:
      <ul>
        <li>Palabras en logo: <em>Codesign, Learning, Itineraries</em></li>
        <li>Lema en cabecera: <strong>Codiseño de aprendizaje flexible</strong></li>
      </ul>
    </li>
    <li><strong>Eje Celeste (Personalización)</strong>:
      <ul>
        <li>Palabras en logo: <em>Enhance, Personalized, Environments</em></li>
        <li>Lema en cabecera: <strong>Itinerarios personalizados y agénticos</strong></li>
      </ul>
    </li>
    <li><strong>Eje Verde (Tecnología)</strong>:
      <ul>
        <li>Palabras en logo: <em>Technology, Inteligencia, Artificial</em></li>
        <li>Lema en cabecera: <strong>Ambientes enriquecidos con Tecnología</strong></li>
      </ul>
    </li>
  </ul>

  <hr>

  <h2>2. Textos de Navegación y Generales</h2>
  <p>Textos estáticos clave utilizados en menús, botones comunes y pie de página en los tres idiomas del sitio:</p>
  
  <table>
    <thead>
      <tr>
        <th>Clave de Traducción (ID)</th>
        <th>Descripción de Uso</th>
        <th>Español (ES)</th>
        <th>Catalán (CA)</th>
        <th>Inglés (EN)</th>
      </tr>
    </thead>
    <tbody>
  """)

  navigation_keys = [
      ("menu_inicio", "Enlace al Inicio en menú"),
      ("menu_proyecto", "Enlace a Proyecto en menú"),
      ("menu_impacto", "Enlace a Impacto en menú"),
      ("search_placeholder", "Buscador de producción científica"),
      ("colab_title", "Cabecera de colaboradores"),
      ("footer_copy", "Derechos de autor en pie de página")
  ]

  for key, label in navigation_keys:
      html.append(f"""      <tr>
        <td><strong>{key}</strong></td>
        <td>{label}</td>
        <td>{translations['es'][key]}</td>
        <td>{translations['ca'][key]}</td>
        <td>{translations['en'][key]}</td>
      </tr>""")

  html.append("""    </tbody>
  </table>

  <hr>

  <h2>3. Vista: Inicio (Home)</h2>
  <p>Textos y secciones correspondientes a la página de bienvenida y presentación general.</p>
  """)

  home_keys = [
      ("hero_tag", "Etiqueta superior del Hero"),
      ("hero_main_title", "Título Principal del Hero"),
      ("hero_desc", "Descripción detallada del Hero"),
      ("btn_conocer", "Texto del Botón Principal"),
      ("progress_label", "Etiqueta de la barra de progreso (Valor fijo: 75%)"),
      ("stats_years", "Estadística 1: Años de Investigación"),
      ("stats_investigadores", "Estadística 2: Investigadores"),
      ("stats_publicaciones", "Estadística 3: Publicaciones"),
      ("stats_experiencias", "Estadística 4: Experiencias en centros"),
      ("news_title", "Cabecera de noticias en Inicio")
  ]

  for key, label in home_keys:
      html.append(f"""  <h3>ID: <code>{key}</code> ({label})</h3>
  <div class="lang-section">
    <div class="lang-tag">Español (ES)</div>
    <div>{translations['es'][key]}</div>
    <div class="lang-tag">Catalán (CA)</div>
    <div>{translations['ca'][key]}</div>
    <div class="lang-tag">Inglés (EN)</div>
    <div>{translations['en'][key]}</div>
  </div>""")

  html.append("""  <hr>

  <h2>4. Vista: Proyecto</h2>
  <p>Información detallada sobre el proyecto de investigación, sus objetivos y su equipo.</p>

  <h3>Pestañas del Submenú Interno</h3>
  <table>
    <thead>
      <tr>
        <th>Pestaña</th>
        <th>Español (ES)</th>
        <th>Catalán (CA)</th>
        <th>Inglés (EN)</th>
      </tr>
    </thead>
    <tbody>
      <tr>
        <td><strong>Descripción (<code>submenu_desc</code>)</strong></td>
        <td>""" + translations['es']['submenu_desc'] + """</td>
        <td>""" + translations['ca']['submenu_desc'] + """</td>
        <td>""" + translations['en']['submenu_desc'] + """</td>
      </tr>
      <tr>
        <td><strong>Objetivos (<code>submenu_obj</code>)</strong></td>
        <td>""" + translations['es']['submenu_obj'] + """</td>
        <td>""" + translations['ca']['submenu_obj'] + """</td>
        <td>""" + translations['en']['submenu_obj'] + """</td>
      </tr>
      <tr>
        <td><strong>Miembros (<code>submenu_miembros</code>)</strong></td>
        <td>""" + translations['es']['submenu_miembros'] + """</td>
        <td>""" + translations['ca']['submenu_miembros'] + """</td>
        <td>""" + translations['en']['submenu_miembros'] + """</td>
      </tr>
    </tbody>
  </table>

  <h3>4.1. Descripción del Proyecto</h3>
  """)

  desc_keys = [
      ("project_pretitle", "Pretítulo de sección"),
      ("project_title", "Título de sección"),
      ("project_what_is", "Pregunta de cabecera"),
      ("project_what_is_p1", "Primer párrafo descriptivo"),
      ("project_what_is_p2", "Segundo párrafo descriptivo"),
      ("project_what_is_p3", "Tercer párrafo descriptivo")
  ]

  for key, label in desc_keys:
      html.append(f"""  <h4>ID: <code>{key}</code> ({label})</h4>
  <div class="lang-section">
    <div class="lang-tag">Español (ES)</div>
    <div>{translations['es'][key]}</div>
    <div class="lang-tag">Catalán (CA)</div>
    <div>{translations['ca'][key]}</div>
    <div class="lang-tag">Inglés (EN)</div>
    <div>{translations['en'][key]}</div>
  </div>""")

  html.append("""  <h4>Ficha Técnica / Metadatos de la Investigación</h4>
  """)

  meta_keys = [
      ("meta_ref", "Ref. Proyecto", "PID2021-127836NB-I00"),
      ("meta_duracion", "Duración", "2022-2025"),
      ("meta_lider", "Institución Líder", "UIB (Universitat de les Illes Balears)"),
      ("meta_financiacion", "Financiación", "MCIN / AEI")
  ]

  for key, label, val in meta_keys:
      html.append(f"""  <div class="meta-box">
    <strong>{label} (<code>{key}</code>)</strong>:<br>
    Etiquetas: ES: "{translations['es'][key]}" | CA: "{translations['ca'][key]}" | EN: "{translations['en'][key]}"<br>
    Valor Fijo: <code>{val}</code>
  </div>""")

  html.append("""  <h4>Tarjeta del Framework Metodológico</h4>
  <div class="meta-box">
    <strong>Título (<code>framework_header</code>)</strong>: """ + translations['es']['framework_header'] + """<br>
    <strong>Subtítulo (<code>framework_subtitle</code>)</strong>: ES: \"""" + translations['es']['framework_subtitle'] + """\" | CA: \"""" + translations['ca']['framework_subtitle'] + """\" | EN: \"""" + translations['en']['framework_subtitle'] + """\"<br>
    <strong>Etiquetas de especialidades (HTML)</strong>: <code>Codiseño</code> | <code>Telecolaboración</code> | <code>IA & NLP</code> | <code>CALL</code> | <code>Personalización</code> | <code>CMC</code>
  </div>

  <h3>4.2. Objetivos del Proyecto</h3>
  <p><strong>Cabecera (<code>obj_title</code>)</strong>: ES: "<i>""" + translations['es']['obj_title'] + """</i>" | CA: "<i>""" + translations['ca']['obj_title'] + """</i>" | EN: "<i>""" + translations['en']['obj_title'] + """</i>"</p>
  """)

  obj_items = [
      ("obj_1_title", "obj_1_desc", "Objetivo 1 (Principal)"),
      ("obj_2_title", "obj_2_desc", "Objetivo 2 (Innovación)"),
      ("obj_3_title", "obj_3_desc", "Objetivo 3 (Alcance Internacional)"),
      ("obj_4_title", "obj_4_desc", "Objetivo 4 (Impacto Social)")
  ]

  for t_key, d_key, label in obj_items:
      html.append(f"""  <h4>{label}</h4>
  <div class="lang-section">
    <strong>Título (<code>{t_key}</code>)</strong>:
    <ul>
      <li>ES: "{translations['es'][t_key]}"</li>
      <li>CA: "{translations['ca'][t_key]}"</li>
      <li>EN: "{translations['en'][t_key]}"</li>
    </ul>
    <strong>Descripción (<code>{d_key}</code>)</strong>:
    <ul>
      <li>ES: "{translations['es'][d_key]}"</li>
      <li>CA: "{translations['ca'][d_key]}"</li>
      <li>EN: "{translations['en'][d_key]}"</li>
    </ul>
  </div>""")

  html.append("""  <h3>4.3. Miembros del Equipo de Investigación</h3>
  <p><strong>Cabecera (<code>tab_equipo</code>)</strong>: ES: "<i>""" + translations['es']['tab_equipo'] + """</i>" | CA: "<i>""" + translations['ca']['tab_equipo'] + """</i>" | EN: "<i>""" + translations['en']['tab_equipo'] + """</i>"</p>
  """)

  for member in team_members:
      html.append(f"""  <div style="border: 1px solid #cbd5e1; padding: 15px; margin-bottom: 20px; border-radius: 4px; background-color: #fdfdfd;">
    <h4 style="margin-top: 0; color: #1e293b;">{member['name']}</h4>
    <div style="font-size: 0.9em; color: #475569; margin-bottom: 10px;">
      <strong>Rol</strong>: ES: "{member['role']['es']}" | CA: "{member['role']['ca']}" | EN: "{member['role']['en']}"<br>
      <strong>Contacto</strong>: Email: <a href="mailto:{member['email']}">{member['email']}</a> | ORCID: {member['orcid']}
    </div>
    
    <strong>Biografías por idioma:</strong>
    <div class="lang-section">
      <div class="lang-tag">Español (ES)</div>
      <div>{member['bio']['es']}</div>
      <div class="lang-tag">Catalán (CA)</div>
      <div>{member['bio']['ca']}</div>
      <div class="lang-tag">Inglés (EN)</div>
      <div>{member['bio']['en']}</div>
    </div>
  </div>""")

  html.append("""  <hr>

  <h2>5. Vista: Impacto y Difusión</h2>
  <p>Actividades del proyecto, producción de artículos en revistas indexadas y recursos didácticos descargables.</p>

  <h3>Pestañas del Submenú Interno</h3>
  <table>
    <thead>
      <tr>
        <th>Pestaña</th>
        <th>Español (ES)</th>
        <th>Catalán (CA)</th>
        <th>Inglés (EN)</th>
      </tr>
    </thead>
    <tbody>
      <tr>
        <td><strong>Últimas Noticias (<code>submenu_noticias</code>)</strong></td>
        <td>""" + translations['es']['submenu_noticias'] + """</td>
        <td>""" + translations['ca']['submenu_noticias'] + """</td>
        <td>""" + translations['en']['submenu_noticias'] + """</td>
      </tr>
      <tr>
        <td><strong>Transferencia (<code>submenu_transferencia</code>)</strong></td>
        <td>""" + translations['es']['submenu_transferencia'] + """</td>
        <td>""" + translations['ca']['submenu_transferencia'] + """</td>
        <td>""" + translations['en']['submenu_transferencia'] + """</td>
      </tr>
      <tr>
        <td><strong>Producción Científica (<code>submenu_publicaciones</code>)</strong></td>
        <td>""" + translations['es']['submenu_publicaciones'] + """</td>
        <td>""" + translations['ca']['submenu_publicaciones'] + """</td>
        <td>""" + translations['en']['submenu_publicaciones'] + """</td>
      </tr>
      <tr>
        <td><strong>Recursos (<code>submenu_recursos</code>)</strong></td>
        <td>""" + translations['es']['submenu_recursos'] + """</td>
        <td>""" + translations['ca']['submenu_recursos'] + """</td>
        <td>""" + translations['en']['submenu_recursos'] + """</td>
      </tr>
    </tbody>
  </table>

  <h3>5.1. Actividades de Transferencia (Detalle Fichas Modales)</h3>
  <p><strong>Título Sección</strong>: """ + translations['es']['trans_title'] + """<br>
  <strong>Subtítulo Sección (<code>trans_subtitle</code>)</strong>:<br>
  ES: \"""" + translations['es']['trans_subtitle'] + """\"<br>
  CA: \"""" + translations['ca']['trans_subtitle'] + """\"<br>
  EN: \"""" + translations['en']['trans_subtitle'] + """\"</p>
  """)

  for act in transfer_activities:
      html.append(f"""  <div style="border: 1px solid #cbd5e1; padding: 15px; margin-bottom: 20px; border-radius: 4px; background-color: #fdfdfd;">
    <h4 style="margin-top: 0; color: #0f766e;"><span class="badge">{act['tag']['es']}</span> {act['title']['es']}</h4>
    <div style="font-size: 0.9em; color: #475569; margin-bottom: 10px;">
      <strong>Fecha</strong>: {act['date']} | <strong>Lugar</strong>: {act['location']}<br>
      <strong>Etiquetas</strong>: {', '.join(act['pills'])}
    </div>
    
    <strong>Títulos y Descripciones Breves:</strong>
    <div class="lang-section">
      <div class="lang-tag">Español (ES)</div>
      <div><strong>Título:</strong> "{act['title']['es']}"<br><strong>Descripción:</strong> {act['desc']['es']}</div>
      <div class="lang-tag">Catalán (CA)</div>
      <div><strong>Título:</strong> "{act['title']['ca']}"<br><strong>Descripción:</strong> {act['desc']['ca']}</div>
      <div class="lang-tag">Inglés (EN)</div>
      <div><strong>Título:</strong> "{act['title']['en']}"<br><strong>Descripción:</strong> {act['desc']['en']}</div>
    </div>
    
    <strong>Texto Extendido (Ficha de detalle en Modal):</strong>
    <div class="lang-section">
      <div class="lang-tag">Español (ES)</div>
      <div style="font-size: 0.95em; text-align: justify; margin-top:5px;">{act['loremIpsum']['es'].replace('\n\n', '<br><br>')}</div>
      <div class="lang-tag">Catalán (CA)</div>
      <div style="font-size: 0.95em; text-align: justify; margin-top:5px;">{act['loremIpsum']['ca'].replace('\n\n', '<br><br>')}</div>
      <div class="lang-tag">Inglés (EN)</div>
      <div style="font-size: 0.95em; text-align: justify; margin-top:5px;">{act['loremIpsum']['en'].replace('\n\n', '<br><br>')}</div>
    </div>
  </div>""")

  html.append("""  <h3>5.2. Producción Científica (Biblioteca Zotero)</h3>
  <p><strong>Título Sección (<code>pub_title</code>)</strong>: ES: \"""" + translations['es']['pub_title'] + """\" | CA: \"""" + translations['ca']['pub_title'] + """\" | EN: \"""" + translations['en']['pub_title'] + """\"<br>
  <strong>Subtítulo Sección (<code>pub_subtitle</code>)</strong>:<br>
  ES: \"""" + translations['es']['pub_subtitle'] + """\"<br>
  CA: \"""" + translations['ca']['pub_subtitle'] + """\"<br>
  EN: \"""" + translations['en']['pub_subtitle'] + """\"</p>
  """)

  for pub in publications:
      doi_isbn_str = ""
      if 'doi' in pub: doi_isbn_str = f"| <strong>DOI:</strong> {pub['doi']}"
      if 'isbn' in pub: doi_isbn_str = f"| <strong>ISBN:</strong> {pub['isbn']}"
      
      html.append(f"""  <div style="border: 1px solid #cbd5e1; padding: 15px; margin-bottom: 20px; border-radius: 4px; background-color: #fdfdfd;">
    <h4 style="margin-top: 0; color: #1e293b;"><span class="badge">{pub['extraLabel']['es']}</span> {pub['title']['es']}</h4>
    <div style="font-size: 0.9em; color: #475569; margin-bottom: 10px;">
      <strong>Referencia:</strong> {pub['citation']}<br>
      <strong>Zotero Key:</strong> <code>{pub['zoteroKey']}</code> {doi_isbn_str} | <strong>Tags:</strong> {', '.join(pub['tags'])}
    </div>
    
    <strong>Títulos y Abstracts por Idioma:</strong>
    <div class="lang-section">
      <div class="lang-tag">Español (ES)</div>
      <div><strong>Título:</strong> "{pub['title']['es']}"<br><strong>Resumen:</strong> {pub['abstract']['es']}</div>
      <div class="lang-tag">Catalán (CA)</div>
      <div><strong>Título:</strong> "{pub['title']['ca']}"<br><strong>Resumen:</strong> {pub['abstract']['ca']}</div>
      <div class="lang-tag">Inglés (EN)</div>
      <div><strong>Título:</strong> "{pub['title']['en']}"<br><strong>Resumen:</strong> {pub['abstract']['en']}</div>
    </div>
  </div>""")

  html.append("""  <h3>5.3. Recursos y Materiales (Descargas)</h3>
  <p><strong>Título Sección (<code>rec_title</code>)</strong>: ES: \"""" + translations['es']['rec_title'] + """\" | CA: \"""" + translations['ca']['rec_title'] + """\" | EN: \"""" + translations['en']['rec_title'] + """\"<br>
  <strong>Subtítulo Sección (<code>rec_subtitle</code>)</strong>:<br>
  ES: \"""" + translations['es']['rec_subtitle'] + """\"<br>
  CA: \"""" + translations['ca']['rec_subtitle'] + """\"<br>
  EN: \"""" + translations['en']['rec_subtitle'] + """\"</p>
  """)

  for res in resources_hardcoded:
      btn_k = res['btn_key']
      html.append(f"""  <div style="border: 1px solid #cbd5e1; padding: 15px; margin-bottom: 15px; border-radius: 4px; background-color: #fdfdfd;">
    <h4 style="margin-top: 0; color: #1e293b;"><span class="badge">{res['category']}</span> {res['title']}</h4>
    <p style="margin: 5px 0 10px 0; font-size: 0.95em;">{res['desc']}</p>
    <div style="font-size: 0.85em; color: #64748b;">
      <strong>Botón Descarga (<code>{btn_k}</code>)</strong>:
      ES: "{translations['es'][btn_k]}" | CA: "{translations['ca'][btn_k]}" | EN: "{translations['en'][btn_k]}"
    </div>
  </div>""")

  html.append("""
</body>
</html>
""")
  return "".join(html)


def generate_docx(output_path="COPLITELE-IA_Secciones_Web.docx"):
    """Tries to generate a DOCX file using python-docx library. Installs it if missing."""
    print("Verificando librería 'python-docx'...")
    try:
        import docx
    except ImportError:
        print("Instalando 'python-docx' para exportación a Word...")
        import subprocess
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            import docx
        except Exception as e:
            print(f"Error al instalar python-docx: {e}")
            print("No se ha podido generar el archivo .docx, pero se han creado los archivos .html y .md.")
            return False

    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    
    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("COPLITELE-IA\n")
    run.font.name = 'Helvetica'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(29, 91, 254) # #1D5BFE
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Secciones de la Web para Verificación de Clientes\n(Documento Multilingüe: ES, CA, EN)\n")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(71, 85, 105)
    
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = intro.add_run(
        "Este documento recopila la estructura completa y el contenido textual de cada sección "
        "de la página web de COPLITELE-IA en sus tres idiomas (Español, Catalán e Inglés). "
        "Su finalidad es facilitar la lectura, revisión y validación de la información con los clientes. "
        "Se divide en secciones principales correspondientes a cada vista de la aplicación Single Page."
    )
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # Section 1: Identidad Visual
    h = doc.add_heading(level=1)
    run = h.add_run("1. Identidad Visual y Conceptos de Marca")
    run.font.color.rgb = RGBColor(29, 91, 254)
    
    p = doc.add_paragraph(
        "El logotipo dinámico de la web incorpora tres ejes de desarrollo en su animación, "
        "cada uno asociado a un color y tres palabras conceptuales en inglés, además de un lema de soporte:"
    )
    
    p_blue = doc.add_paragraph(style='List Bullet')
    p_blue.add_run("Eje Azul (Codiseño):").bold = True
    p_blue.add_run("\n  - Palabras en logo: Codesign, Learning, Itineraries\n  - Lema en cabecera: ")
    p_blue.add_run("Codiseño de aprendizaje flexible").bold = True
    
    p_teal = doc.add_paragraph(style='List Bullet')
    p_teal.add_run("Eje Celeste (Personalización):").bold = True
    p_teal.add_run("\n  - Palabras en logo: Enhance, Personalized, Environments\n  - Lema en cabecera: ")
    p_teal.add_run("Itinerarios personalizados y agénticos").bold = True
    
    p_green = doc.add_paragraph(style='List Bullet')
    p_green.add_run("Eje Verde (Tecnología):").bold = True
    p_green.add_run("\n  - Palabras en logo: Technology, Inteligencia, Artificial\n  - Lema en cabecera: ")
    p_green.add_run("Ambientes enriquecidos con Tecnología").bold = True

    # Section 2: Navegación y Generales
    h = doc.add_heading(level=1)
    run = h.add_run("2. Textos de Navegación y Generales")
    run.font.color.rgb = RGBColor(29, 91, 254)
    
    navigation_keys = [
        ("menu_inicio", "Enlace a Inicio en menú"),
        ("menu_proyecto", "Enlace a Proyecto en menú"),
        ("menu_impacto", "Enlace a Impacto en menú"),
        ("search_placeholder", "Buscador de producción científica"),
        ("colab_title", "Cabecera de colaboradores"),
        ("footer_copy", "Derechos de autor en pie de página")
    ]
    
    for key, label in navigation_keys:
        p = doc.add_paragraph()
        run = p.add_run(f"ID: {key} ({label})")
        run.bold = True
        run.font.size = Pt(12)
        p.add_run(f"\n- Español (ES): {translations['es'][key]}")
        p.add_run(f"\n- Catalán (CA): {translations['ca'][key]}")
        p.add_run(f"\n- Inglés (EN): {translations['en'][key]}")
        
    # Section 3: Inicio
    h = doc.add_heading(level=1)
    run = h.add_run("3. Vista: Inicio (Home)")
    run.font.color.rgb = RGBColor(29, 91, 254)
    
    home_keys = [
        ("hero_tag", "Etiqueta superior del Hero"),
        ("hero_main_title", "Título Principal del Hero"),
        ("hero_desc", "Descripción detallada del Hero"),
        ("btn_conocer", "Texto del Botón Principal"),
        ("progress_label", "Etiqueta de la barra de progreso (Valor fijo: 75%)"),
        ("stats_years", "Estadística 1: Años de Investigación"),
        ("stats_investigadores", "Estadística 2: Investigadores"),
        ("stats_publicaciones", "Estadística 3: Publicaciones"),
        ("stats_experiencias", "Estadística 4: Experiencias en centros"),
        ("news_title", "Cabecera de noticias en Inicio")
    ]
    
    for key, label in home_keys:
        p = doc.add_paragraph()
        run = p.add_run(f"ID: {key} ({label})")
        run.bold = True
        p.add_run(f"\n- Español (ES): {translations['es'][key].replace('<strong>', '').replace('</strong>', '')}")
        p.add_run(f"\n- Catalán (CA): {translations['ca'][key].replace('<strong>', '').replace('</strong>', '')}")
        p.add_run(f"\n- Inglés (EN): {translations['en'][key].replace('<strong>', '').replace('</strong>', '')}")
        
    # Section 4: Proyecto
    h = doc.add_heading(level=1)
    run = h.add_run("4. Vista: Proyecto")
    run.font.color.rgb = RGBColor(29, 91, 254)
    
    p = doc.add_paragraph()
    run = p.add_run("Pestañas del Submenú Interno")
    run.bold = True
    p.add_run(f"\n- Descripción (submenu_desc): ES: \"{translations['es']['submenu_desc']}\" | CA: \"{translations['ca']['submenu_desc']}\" | EN: \"{translations['en']['submenu_desc']}\"")
    p.add_run(f"\n- Objetivos (submenu_obj): ES: \"{translations['es']['submenu_obj']}\" | CA: \"{translations['ca']['submenu_obj']}\" | EN: \"{translations['en']['submenu_obj']}\"")
    p.add_run(f"\n- Miembros (submenu_miembros): ES: \"{translations['es']['submenu_miembros']}\" | CA: \"{translations['ca']['submenu_miembros']}\" | EN: \"{translations['en']['submenu_miembros']}\"")
    
    doc.add_heading("4.1. Descripción del Proyecto", level=2)
    
    desc_keys = [
        ("project_pretitle", "Pretítulo de sección"),
        ("project_title", "Título de sección"),
        ("project_what_is", "Pregunta de cabecera"),
        ("project_what_is_p1", "Primer párrafo descriptivo"),
        ("project_what_is_p2", "Segundo párrafo descriptivo"),
        ("project_what_is_p3", "Tercer párrafo descriptivo")
    ]
    
    for key, label in desc_keys:
        p = doc.add_paragraph()
        run = p.add_run(f"ID: {key} ({label})")
        run.bold = True
        p.add_run(f"\n- Español (ES): {translations['es'][key].replace('<strong>', '').replace('</strong>', '')}")
        p.add_run(f"\n- Catalán (CA): {translations['ca'][key].replace('<strong>', '').replace('</strong>', '')}")
        p.add_run(f"\n- Inglés (EN): {translations['en'][key].replace('<strong>', '').replace('</strong>', '')}")
        
    doc.add_heading("Ficha Técnica / Metadatos de la Investigación", level=3)
    meta_keys = [
        ("meta_ref", "Ref. Proyecto", "PID2021-127836NB-I00"),
        ("meta_duracion", "Duración", "2022-2025"),
        ("meta_lider", "Institución Líder", "UIB (Universitat de les Illes Balears)"),
        ("meta_financiacion", "Financiación", "MCIN / AEI")
    ]
    for key, label, val in meta_keys:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} ({key})")
        run.bold = True
        p.add_run(f"\n- Etiquetas: ES: \"{translations['es'][key]}\" | CA: \"{translations['ca'][key]}\" | EN: \"{translations['en'][key]}\"")
        p.add_run(f"\n- Valor Fijo: {val}")
        
    doc.add_heading("Tarjeta del Framework Metodológico", level=3)
    p = doc.add_paragraph()
    p.add_run("Título (framework_header): ").bold = True
    p.add_run(f"{translations['es']['framework_header']}\n")
    p.add_run("Subtítulo (framework_subtitle): ").bold = True
    p.add_run(f"ES: \"{translations['es']['framework_subtitle']}\" | CA: \"{translations['ca']['framework_subtitle']}\" | EN: \"{translations['en']['framework_subtitle']}\"\n")
    p.add_run("Etiquetas fijas: ").bold = True
    p.add_run("Codiseño, Telecolaboración, IA & NLP, CALL, Personalización, CMC")
    
    doc.add_heading("4.2. Objetivos del Proyecto", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Cabecera (obj_title): ").bold = True
    p.add_run(f"ES: \"{translations['es']['obj_title']}\" | CA: \"{translations['ca']['obj_title']}\" | EN: \"{translations['en']['obj_title']}\"")
    
    obj_items = [
        ("obj_1_title", "obj_1_desc", "Objetivo 1 (Principal)"),
        ("obj_2_title", "obj_2_desc", "Objetivo 2 (Innovación)"),
        ("obj_3_title", "obj_3_desc", "Objetivo 3 (Alcance Internacional)"),
        ("obj_4_title", "obj_4_desc", "Objetivo 4 (Impacto Social)")
    ]
    
    for t_key, d_key, label in obj_items:
        doc.add_heading(label, level=3)
        p = doc.add_paragraph()
        p.add_run("Título: ").bold = True
        p.add_run(f"ES: \"{translations['es'][t_key]}\" | CA: \"{translations['ca'][t_key]}\" | EN: \"{translations['en'][t_key]}\"\n")
        p.add_run("Descripción: ").bold = True
        p.add_run(f"\n- ES: \"{translations['es'][d_key]}\"")
        p.add_run(f"\n- CA: \"{translations['ca'][d_key]}\"")
        p.add_run(f"\n- EN: \"{translations['en'][d_key]}\"")
        
    doc.add_heading("4.3. Miembros del Equipo de Investigación", level=2)
    p = doc.add_paragraph()
    p.add_run("Cabecera (tab_equipo): ").bold = True
    p.add_run(f"ES: \"{translations['es']['tab_equipo']}\" | CA: \"{translations['ca']['tab_equipo']}\" | EN: \"{translations['en']['tab_equipo']}\"")
    
    for member in team_members:
        doc.add_heading(member['name'], level=3)
        p = doc.add_paragraph()
        p.add_run("Rol: ").bold = True
        p.add_run(f"ES: \"{member['role']['es']}\" | CA: \"{member['role']['ca']}\" | EN: \"{member['role']['en']}\"\n")
        p.add_run("Contacto: ").bold = True
        p.add_run(f"Email: {member['email']} | ORCID: {member['orcid']}")
        
        p = doc.add_paragraph()
        p.add_run("Biografías por idioma:").bold = True
        p.add_run(f"\n- ES: {member['bio']['es']}")
        p.add_run(f"\n- CA: {member['bio']['ca']}")
        p.add_run(f"\n- EN: {member['bio']['en']}")
        
    # Section 5: Impacto
    h = doc.add_heading(level=1)
    run = h.add_run("5. Vista: Impacto y Difusión")
    run.font.color.rgb = RGBColor(29, 91, 254)
    
    p = doc.add_paragraph()
    run = p.add_run("Pestañas del Submenú Interno")
    run.bold = True
    p.add_run(f"\n- Noticias (submenu_noticias): ES: \"{translations['es']['submenu_noticias']}\" | CA: \"{translations['ca']['submenu_noticias']}\" | EN: \"{translations['en']['submenu_noticias']}\"")
    p.add_run(f"\n- Transferencia (submenu_transferencia): ES: \"{translations['es']['submenu_transferencia']}\" | CA: \"{translations['ca']['submenu_transferencia']}\" | EN: \"{translations['en']['submenu_transferencia']}\"")
    p.add_run(f"\n- Producción Científica (submenu_publicaciones): ES: \"{translations['es']['submenu_publicaciones']}\" | CA: \"{translations['ca']['submenu_publicaciones']}\" | EN: \"{translations['en']['submenu_publicaciones']}\"")
    p.add_run(f"\n- Recursos (submenu_recursos): ES: \"{translations['es']['submenu_recursos']}\" | CA: \"{translations['ca']['submenu_recursos']}\" | EN: \"{translations['en']['submenu_recursos']}\"")
    
    doc.add_heading("5.1. Actividades de Transferencia (Fichas de detalle)", level=2)
    p = doc.add_paragraph()
    p.add_run("Título de la sección: ").bold = True
    p.add_run(f"{translations['es']['trans_title']}\n")
    p.add_run("Subtítulo de la sección (trans_subtitle): ").bold = True
    p.add_run(f"ES: \"{translations['es']['trans_subtitle']}\" | CA: \"{translations['ca']['trans_subtitle']}\" | EN: \"{translations['en']['trans_subtitle']}\"")
    
    for act in transfer_activities:
        doc.add_heading(f"[{act['tag']['es']}] {act['title']['es']}", level=3)
        p = doc.add_paragraph()
        p.add_run("Fecha: ").bold = True
        p.add_run(f"{act['date']} | ")
        p.add_run("Ubicación: ").bold = True
        p.add_run(f"{act['location']} | ")
        p.add_run("Etiquetas: ").bold = True
        p.add_run(f"{', '.join(act['pills'])}\n")
        
        p = doc.add_paragraph()
        p.add_run("Títulos y Descripciones Cortas:").bold = True
        p.add_run(f"\n- ES: Título: \"{act['title']['es']}\" | Descripción: {act['desc']['es']}")
        p.add_run(f"\n- CA: Título: \"{act['title']['ca']}\" | Descripción: {act['desc']['ca']}")
        p.add_run(f"\n- EN: Título: \"{act['title']['en']}\" | Descripción: {act['desc']['en']}")
        
        p = doc.add_paragraph()
        p.add_run("Texto Detallado Completo (Modal):").bold = True
        p.add_run(f"\n- ES: {act['loremIpsum']['es']}")
        p.add_run(f"\n- CA: {act['loremIpsum']['ca']}")
        p.add_run(f"\n- EN: {act['loremIpsum']['en']}")
        
    doc.add_heading("5.2. Producción Científica (Biblioteca Zotero)", level=2)
    p = doc.add_paragraph()
    p.add_run("Título (pub_title): ").bold = True
    p.add_run(f"ES: \"{translations['es']['pub_title']}\" | CA: \"{translations['ca']['pub_title']}\" | EN: \"{translations['en']['pub_title']}\"\n")
    p.add_run("Subtítulo (pub_subtitle): ").bold = True
    p.add_run(f"ES: \"{translations['es']['pub_subtitle']}\" | CA: \"{translations['ca']['pub_subtitle']}\" | EN: \"{translations['en']['pub_subtitle']}\"")
    
    for pub in publications:
        doc.add_heading(f"[{pub['extraLabel']['es']}] {pub['title']['es']}", level=3)
        p = doc.add_paragraph()
        p.add_run("Cita: ").bold = True
        p.add_run(f"{pub['citation']} | ")
        p.add_run("Zotero Key: ").bold = True
        p.add_run(f"{pub['zoteroKey']} | ")
        if 'doi' in pub:
            p.add_run(f"DOI: {pub['doi']} | ")
        if 'isbn' in pub:
            p.add_run(f"ISBN: {pub['isbn']} | ")
        p.add_run("Tags: ").bold = True
        p.add_run(f"{', '.join(pub['tags'])}\n")
        
        p = doc.add_paragraph()
        p.add_run("Títulos y Abstracts por Idioma:").bold = True
        p.add_run(f"\n- ES: Título: \"{pub['title']['es']}\" | Resumen: {pub['abstract']['es']}")
        p.add_run(f"\n- CA: Título: \"{pub['title']['ca']}\" | Resumen: {pub['abstract']['ca']}")
        p.add_run(f"\n- EN: Título: \"{pub['title']['en']}\" | Resumen: {pub['abstract']['en']}")
        
    doc.add_heading("5.3. Recursos y Materiales (Descargables)", level=2)
    p = doc.add_paragraph()
    p.add_run("Título (rec_title): ").bold = True
    p.add_run(f"ES: \"{translations['es']['rec_title']}\" | CA: \"{translations['ca']['rec_title']}\" | EN: \"{translations['en']['rec_title']}\"\n")
    p.add_run("Subtítulo (rec_subtitle): ").bold = True
    p.add_run(f"ES: \"{translations['es']['rec_subtitle']}\" | CA: \"{translations['ca']['rec_subtitle']}\" | EN: \"{translations['en']['rec_subtitle']}\"")
    
    for res in resources_hardcoded:
        doc.add_heading(f"[{res['category']}] {res['title']}", level=3)
        p = doc.add_paragraph()
        p.add_run("Descripción (ES): ").bold = True
        p.add_run(f"{res['desc']}\n")
        btn_k = res['btn_key']
        p.add_run("Botón de Descarga: ").bold = True
        p.add_run(f"ES: \"{translations['es'][btn_k]}\" | CA: \"{translations['ca'][btn_k]}\" | EN: \"{translations['en'][btn_k]}\"")
        
    doc.save(output_path)
    print(f"Archivo de Word creado con éxito: {output_path}")
    return True


if __name__ == "__main__":
    print("Iniciando generación de documentos...")
    
    # Ensure public directories exist
    os.makedirs("public/COPLITELE-IA_Secciones_Web", exist_ok=True)
    
    # Generate content
    md_content = generate_markdown()
    html_content = generate_html()
    
    # Write to local root (for convenience)
    with open("COPLITELE-IA_Secciones_Web.md", "w", encoding="utf-8") as f:
        f.write(md_content)
    with open("COPLITELE-IA_Secciones_Web.html", "w", encoding="utf-8") as f:
        f.write(html_content)
        
    # Write to public/ for deployment (with standard filenames)
    with open("public/COPLITELE-IA_Secciones_Web.md", "w", encoding="utf-8") as f:
        f.write(md_content)
    with open("public/COPLITELE-IA_Secciones_Web.html", "w", encoding="utf-8") as f:
        f.write(html_content)
        
    # Write to public/COPLITELE-IA_Secciones_Web/ for routing index.html
    with open("public/COPLITELE-IA_Secciones_Web/index.html", "w", encoding="utf-8") as f:
        f.write(html_content)
    with open("public/COPLITELE-IA_Secciones_Web/COPLITELE-IA_Secciones_Web.md", "w", encoding="utf-8") as f:
        f.write(md_content)
        
    # Generate DOCX files
    docx_success_root = generate_docx("COPLITELE-IA_Secciones_Web.docx")
    docx_success_public = generate_docx("public/COPLITELE-IA_Secciones_Web.docx")
    docx_success_folder = generate_docx("public/COPLITELE-IA_Secciones_Web/COPLITELE-IA_Secciones_Web.docx")
    
    print("\n¡Proceso finalizado con éxito!")
    if docx_success_root and docx_success_public and docx_success_folder:
        print("Se han creado y copiado los archivos en la carpeta raíz y en la carpeta 'public' (para despliegue en GitHub Pages).")
    else:
        print("Se han creado los archivos HTML y MD, pero falló la generación de DOCX.")
